"""
Document Processor for OrientaTech RAG System
Extracts text from various document formats (PDF, DOC, DOCX)
"""
import os
import re
from typing import Dict, Optional
from pathlib import Path
from loguru import logger
import fitz  # PyMuPDF
from docx import Document


class DocumentProcessor:
    """
    Service for extracting text from document files
    
    Supported formats:
    - PDF (.pdf) - Using PyMuPDF
    - Word Document (.docx) - Using python-docx
    - Text (.txt) - Direct reading
    
    Features:
    - Auto-detection of file format
    - Text cleaning and normalization
    - Metadata extraction
    - Error handling
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.txt']
    
    def __init__(self):
        """Initialize the document processor"""
        logger.debug("DocumentProcessor initialized")
    
    # ===================================
    # UNIVERSAL METHOD (Main API)
    # ===================================
    
    def extract_text(self, file_path: Path) -> str:
        """
        Universal method to extract text from any supported document
        Automatically detects file format and uses appropriate method
        
        Args:
            file_path: Path to the document file
            
        Returns:
            str: Extracted and cleaned text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            RuntimeError: If extraction fails
            
        Usage:
            processor = DocumentProcessor()
            text = processor.extract_text("resume.pdf")  # Auto-detects PDF
            text = processor.extract_text("letter.docx")  # Auto-detects DOCX
        """
        # Convert to Path object if string
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check if format is supported
        if not self.is_supported_format(file_path):
            raise ValueError(
                f"Unsupported file format: {file_path.suffix}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        try:
            # Route to appropriate extraction method
            if extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {extension}")
            
            # Clean extracted text
            cleaned_text = self.clean_extracted_text(text)
            
            logger.info(f"✅ Extracted {len(cleaned_text)} characters from {file_path.name}")
            return cleaned_text
            
        except Exception as e:
            logger.error(f"❌ Error extracting text from {file_path.name}: {e}")
            raise RuntimeError(f"Failed to extract text: {e}")
    
    # ===================================
    # SPECIFIC METHODS (Advanced API)
    # ===================================
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using PyMuPDF
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text from all pages
        """
        try:
            text_content = []
            
            # Open PDF
            with fitz.open(file_path) as pdf_document:
                # Extract text from each page
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    text = page.get_text()
                    text_content.append(text)
            
            # Join all pages
            full_text = "\n\n".join(text_content)
            
            logger.debug(f"Extracted text from {len(pdf_document)} pages")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            raise RuntimeError(f"PDF extraction failed: {e}")
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text from all paragraphs
        """
        try:
            # Open DOCX
            document = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Join all paragraphs
            full_text = "\n\n".join(paragraphs)
            
            logger.debug(f"Extracted text from {len(paragraphs)} paragraphs")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            raise RuntimeError(f"DOCX extraction failed: {e}")
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            str: File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.debug(f"Read {len(text)} characters from TXT file")
            return text
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                logger.debug("Read TXT file with latin-1 encoding")
                return text
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                raise RuntimeError(f"TXT reading failed: {e}")
        except Exception as e:
            logger.error(f"Error extracting from TXT: {e}")
            raise RuntimeError(f"TXT extraction failed: {e}")
    
    # ===================================
    # TEXT PROCESSING
    # ===================================
    
    def clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove special characters that may cause issues
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    # ===================================
    # METADATA EXTRACTION
    # ===================================
    
    def extract_metadata(self, file_path: Path) -> Dict:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dict: Document metadata
        """
        file_path = Path(file_path)
        
        metadata = {
            'filename': file_path.name,
            'extension': file_path.suffix.lower(),
            'size_bytes': file_path.stat().st_size if file_path.exists() else 0,
            'size_mb': round(file_path.stat().st_size / (1024 * 1024), 2) if file_path.exists() else 0
        }
        
        # Format-specific metadata
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif extension == '.docx':
                metadata.update(self._extract_docx_metadata(file_path))
        except Exception as e:
            logger.warning(f"Could not extract metadata: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict:
        """Extract PDF-specific metadata"""
        try:
            with fitz.open(file_path) as pdf:
                return {
                    'page_count': len(pdf),
                    'author': pdf.metadata.get('author', ''),
                    'title': pdf.metadata.get('title', ''),
                    'creator': pdf.metadata.get('creator', '')
                }
        except Exception as e:
            logger.debug(f"Could not extract PDF metadata: {e}")
            return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict:
        """Extract DOCX-specific metadata"""
        try:
            document = Document(file_path)
            props = document.core_properties
            return {
                'author': props.author or '',
                'title': props.title or '',
                'subject': props.subject or '',
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None
            }
        except Exception as e:
            logger.debug(f"Could not extract DOCX metadata: {e}")
            return {}
    
    # ===================================
    # VALIDATION
    # ===================================
    
    def is_supported_format(self, file_path: Path) -> bool:
        """
        Check if file format is supported
        
        Args:
            file_path: Path to file
            
        Returns:
            bool: True if format is supported
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return self.SUPPORTED_EXTENSIONS.copy()
    
    # ===================================
    # ASYNC WRAPPERS
    # ===================================
    
    async def extract_text_async(self, file_path: Path) -> str:
        """
        Async wrapper for extract_text
        Useful for FastAPI async endpoints
        
        Args:
            file_path: Path to document file
            
        Returns:
            str: Extracted text
        """
        import asyncio
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.extract_text,
            file_path
        )
    
    async def extract_metadata_async(self, file_path: Path) -> Dict:
        """
        Async wrapper for extract_metadata
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dict: Document metadata
        """
        import asyncio
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.extract_metadata,
            file_path
        )


# ===================================
# Singleton instance
# ===================================

_document_processor_instance: Optional[DocumentProcessor] = None

def get_document_processor() -> DocumentProcessor:
    """
    Get singleton instance of DocumentProcessor
    
    Returns:
        DocumentProcessor: Singleton instance
    """
    global _document_processor_instance
    
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    
    return _document_processor_instance


# ===================================
# Convenience functions
# ===================================

def extract_text(file_path: Path) -> str:
    """
    Convenience function to extract text
    
    Args:
        file_path: Path to document file
        
    Returns:
        str: Extracted text
    """
    processor = get_document_processor()
    return processor.extract_text(file_path)


def extract_metadata(file_path: Path) -> Dict:
    """
    Convenience function to extract metadata
    
    Args:
        file_path: Path to document file
        
    Returns:
        Dict: Document metadata
    """
    processor = get_document_processor()
    return processor.extract_metadata(file_path)


# ===================================
# Testing and validation
# ===================================

if __name__ == "__main__":
    """Test the document processor"""
    logger.info("Testing DocumentProcessor...")
    
    # Create processor
    processor = DocumentProcessor()
    
    # Test format validation
    logger.info("Test 1: Format validation")
    assert processor.is_supported_format(Path("test.pdf")) == True
    assert processor.is_supported_format(Path("test.docx")) == True
    assert processor.is_supported_format(Path("test.xlsx")) == False
    logger.success("✅ Format validation works")
    
    # Test supported formats
    logger.info("\nTest 2: Supported formats")
    formats = processor.get_supported_formats()
    logger.info(f"✅ Supported formats: {formats}")
    
    logger.success("🎉 All validation tests passed!")
    logger.info("\n⚠️ Note: File extraction tests require actual document files")

